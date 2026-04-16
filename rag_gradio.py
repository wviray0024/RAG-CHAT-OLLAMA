"""
RAG Gradio App — LangChain + Ollama + PDF/DOCX
Fixed: better PDF extraction, larger chunks, more retrieval coverage
--------------------------------------------------------------------
Requirements:
    pip install gradio langchain langchain-community langchain-ollama
                langchain-text-splitters langchain-chroma langchain-core
                chromadb pdfplumber pymupdf python-docx docx2txt

Run:
    python rag_gradio.py
Then open: http://localhost:7860
"""

import gc
import os
import time
import shutil
import gradio as gr
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_ollama import OllamaEmbeddings, ChatOllama
from langchain_core.prompts import PromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser


# ── Config ────────────────────────────────────────────────────────────────────

OLLAMA_MODEL  = "llama3"
EMBED_MODEL   = "nomic-embed-text"
CHROMA_DIR    = "./chroma_db"

# ── KEY FIX: larger chunks = more context per retrieval ──────────────────────
CHUNK_SIZE    = 1500    # was 500 — too small, cut sentences mid-way
CHUNK_OVERLAP = 200     # was 100 — more overlap = fewer missed details
RETRIEVER_K   = 10      # was 6 — retrieve more chunks for better coverage

ALLOWED_EXTS  = {".pdf", ".docx"}

PROMPT_TEMPLATE = """### SYSTEM:
You are a strict document assistant. Answer the question using ONLY the provided context.

### CONSTRAINTS:
1. Use ONLY the information in the Context section below.
2. If the answer is not in the context, say: "I am sorry, but the uploaded documents do not contain information regarding this request."
3. Do NOT use outside knowledge.
4. Do NOT say "According to the context..." in your answer.
5. Be thorough — include ALL relevant details from the context.

### Context:
{context}

### Question:
{question}

### Answer:"""


# ── Global state ──────────────────────────────────────────────────────────────

vectorstore = None
chain       = None
doc_store   = {}


# ── PDF loader with fallback ───────────────────────────────────────────────────

def load_pdf_best(path: str):
    """
    Try pdfplumber first (best for tables/structured PDFs),
    fall back to PyMuPDF, then PyPDFLoader as last resort.
    """
    pages = []

    # Attempt 1: pdfplumber — best for tables and structured content
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and text.strip():
                    pages.append(Document(
                        page_content=text,
                        metadata={"page": i, "filename": os.path.basename(path), "filetype": ".pdf"}
                    ))
        if pages:
            print(f"   → pdfplumber: extracted {len(pages)} pages")
            return pages
    except Exception as e:
        print(f"   pdfplumber failed: {e}")

    # Attempt 2: PyMuPDF (fitz) — good for complex layouts
    try:
        import fitz
        doc = fitz.open(path)
        for i, page in enumerate(doc):
            text = page.get_text()
            if text and text.strip():
                pages.append(Document(
                    page_content=text,
                    metadata={"page": i, "filename": os.path.basename(path), "filetype": ".pdf"}
                ))
        doc.close()
        if pages:
            print(f"   → PyMuPDF: extracted {len(pages)} pages")
            return pages
    except Exception as e:
        print(f"   PyMuPDF failed: {e}")

    # Attempt 3: PyPDFLoader — basic fallback
    try:
        from langchain_community.document_loaders import PyPDFLoader
        loader = PyPDFLoader(path)
        pages  = loader.load()
        for page in pages:
            page.metadata["filename"] = os.path.basename(path)
            page.metadata["filetype"] = ".pdf"
        print(f"   → PyPDFLoader fallback: extracted {len(pages)} pages")
        return pages
    except Exception as e:
        raise ValueError(f"All PDF loaders failed for {path}: {e}")


def load_document(path: str):
    ext = os.path.splitext(path)[-1].lower()

    if ext == ".pdf":
        pages = load_pdf_best(path)
    elif ext == ".docx":
        docx    = DocxDocument(path)
        content = "\n".join([p.text for p in docx.paragraphs if p.text.strip()])
        pages   = [Document(
            page_content=content,
            metadata={"filename": os.path.basename(path), "page": 0, "filetype": ".docx"},
        )]
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    for i, page in enumerate(pages):
        page.metadata.setdefault("page", i)
        page.metadata["filename"] = os.path.basename(path)
        page.metadata["filetype"] = ext

    return pages


# ── Vector store ──────────────────────────────────────────────────────────────

def rebuild_vectorstore():
    global vectorstore, chain

    chain = None
    if vectorstore is not None:
        try:
            vectorstore._client.close()
        except Exception:
            pass
        vectorstore = None

    gc.collect()
    time.sleep(1)

    if os.path.exists(CHROMA_DIR):
        for attempt in range(5):
            try:
                shutil.rmtree(CHROMA_DIR)
                print(f"Cleared old vector store on attempt {attempt + 1}")
                break
            except PermissionError:
                time.sleep(1)
        else:
            raise Exception(
                "Windows is still locking the database files. "
                "Please restart the app and try again."
            )

    all_pages = []
    for pages in doc_store.values():
        all_pages.extend(pages)

    if not all_pages:
        return

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
    )
    chunks = splitter.split_documents(all_pages)
    print(f"   → Split into {len(chunks)} chunks (size={CHUNK_SIZE}, overlap={CHUNK_OVERLAP})")

    embeddings  = OllamaEmbeddings(model=EMBED_MODEL)
    vectorstore = Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        persist_directory=CHROMA_DIR,
    )
    chain = build_chain(vectorstore)
    print(f"✅ Vector store ready with {len(chunks)} chunks from {len(doc_store)} file(s)")


def build_chain(vs):
    prompt    = PromptTemplate.from_template(PROMPT_TEMPLATE)
    llm       = ChatOllama(model=OLLAMA_MODEL, temperature=0)

    # MMR retrieval = diverse results, reduces repetitive chunks
    retriever = vs.as_retriever(
        search_type="mmr",
        search_kwargs={
            "k": RETRIEVER_K,
            "fetch_k": RETRIEVER_K * 3,   # fetch 3x then pick diverse top-k
            "lambda_mult": 0.7,            # 0=max diversity, 1=max relevance
        },
    )

    def format_docs(docs):
        formatted = []
        for d in docs:
            fname = d.metadata.get("filename", "Unknown")
            page  = d.metadata.get("page", "?")
            formatted.append(f"[Source: {fname} | Page: {page}]\n{d.page_content}")
        return "\n\n---\n\n".join(formatted)

    return (
        {"context": retriever | format_docs, "question": RunnablePassthrough()}
        | prompt
        | llm
        | StrOutputParser()
    )


# ── Helpers ───────────────────────────────────────────────────────────────────

def validate_files(files):
    if not files:
        return [], [("none", "no file")]
    invalid = []
    valid   = []
    for file in files:
        fname = os.path.basename(file.name)
        ext   = os.path.splitext(fname)[-1].lower()
        if ext not in ALLOWED_EXTS:
            invalid.append((fname, ext if ext else "no extension"))
        else:
            valid.append(file)
    return valid, invalid


def format_doc_list():
    if not doc_store:
        return "No documents loaded."
    lines = []
    for fname, pages in doc_store.items():
        ext = os.path.splitext(fname)[-1].upper().replace(".", "")
        lines.append(f"• {fname}  [{ext} — {len(pages)} page(s)]")
    return "\n".join(lines)


def get_doc_choices():
    return list(doc_store.keys())


# ── Gradio handlers ───────────────────────────────────────────────────────────

def upload_files(files):
    global doc_store

    if os.path.exists(CHROMA_DIR) and not doc_store:
        shutil.rmtree(CHROMA_DIR, ignore_errors=True)

    if not files:
        raise gr.Error("No files selected. Please upload at least one .pdf or .docx file.")

    valid_files, invalid = validate_files(files)

    if invalid:
        rejected = ", ".join(f"{fname} ('{ext}')" for fname, ext in invalid)
        raise gr.Error(
            f"Unsupported file type(s): {rejected}. "
            f"Only .pdf and .docx files are accepted."
        )

    newly_added = []
    load_errors = []

    for file in valid_files:
        fname = os.path.basename(file.name)
        if fname in doc_store:
            load_errors.append(f"⚠️ '{fname}' already loaded — skipped.")
            continue
        try:
            pages            = load_document(file.name)
            doc_store[fname] = pages
            newly_added.append(fname)
            print(f"Loaded '{fname}': {len(pages)} page(s)")
        except Exception as e:
            load_errors.append(f"❌ Failed to load '{fname}': {e}")

    if newly_added:
        try:
            rebuild_vectorstore()
        except Exception as e:
            raise gr.Error(f"Error building vector store: {e}")

    parts = []
    if newly_added:
        parts.append("✅ Successfully added:\n" + "\n".join(f"  • {f}" for f in newly_added))
    if load_errors:
        parts.append("\n".join(load_errors))
    parts.append(f"\n📦 Total loaded: {len(doc_store)} file(s)")

    return (
        "\n".join(parts),
        format_doc_list(),
        gr.update(choices=get_doc_choices(), value=None),
    )


def delete_document(filename):
    global doc_store

    if not filename:
        raise gr.Error("Please select a document from the dropdown first.")
    if filename not in doc_store:
        raise gr.Error(f"'{filename}' was not found in loaded documents.")

    del doc_store[filename]
    rebuild_vectorstore()

    status = f"🗑️ Deleted: {filename}\n"
    status += (
        f"📦 Remaining: {len(doc_store)} file(s)"
        if doc_store
        else "📭 No documents loaded. Please upload new files."
    )

    return (
        status,
        format_doc_list(),
        gr.update(choices=get_doc_choices(), value=None),
    )


def delete_all_documents():
    global doc_store, vectorstore, chain

    chain = None
    if vectorstore is not None:
        try:
            vectorstore._client.close()
        except Exception:
            pass
    vectorstore = None
    doc_store   = {}

    gc.collect()
    time.sleep(1)

    if os.path.exists(CHROMA_DIR):
        for _ in range(5):
            try:
                shutil.rmtree(CHROMA_DIR)
                break
            except PermissionError:
                time.sleep(1)

    return (
        "🗑️ All documents deleted. Please upload new files.",
        format_doc_list(),
        gr.update(choices=[], value=None),
        [],
        "",
    )


def hard_reset():
    global vectorstore, chain, doc_store

    chain = None
    if vectorstore is not None:
        try:
            vectorstore._client.close()
        except Exception:
            pass
    vectorstore = None
    doc_store   = {}

    gc.collect()
    time.sleep(1)

    if os.path.exists(CHROMA_DIR):
        for _ in range(5):
            try:
                shutil.rmtree(CHROMA_DIR)
                break
            except PermissionError:
                time.sleep(1)

    return (
        "🔄 Hard reset complete — all documents and vectors cleared.",
        "No documents loaded.",
        gr.update(choices=[], value=None),
        [],
        "",
    )


def chat(message, history):
    global chain

    history = history or []

    if not message.strip():
        return "", history

    if chain is None:
        history.append({"role": "user",      "content": message})
        history.append({"role": "assistant", "content": "⚠️ No documents loaded. Please upload a .pdf or .docx file first."})
        return "", history

    try:
        answer = chain.invoke(message)
        history.append({"role": "user",      "content": message})
        history.append({"role": "assistant", "content": answer})
    except Exception as e:
        history.append({"role": "user",      "content": message})
        history.append({"role": "assistant", "content": f"⚠️ Error: {str(e)}. Please try again."})

    return "", history


def clear_chat():
    return [], ""


# ── Gradio UI ─────────────────────────────────────────────────────────────────

with gr.Blocks(title="RAG Chat — Ollama") as demo:

    gr.Markdown("# Documents Summarizer")
    gr.Markdown("Upload PDF or DOCX files, manage your document library, then chat.")

    with gr.Row():

        # ── Left panel ──────────────────────────────────────────────────────
        with gr.Column(scale=1):

            gr.Markdown("### Upload documents")
            gr.Markdown(
                "<small style='color:gray'>Accepted: <b>.pdf</b> and <b>.docx</b> only.</small>"
            )
            file_input = gr.File(
                label="PDF / DOCX files",
                file_types=[".pdf", ".docx"],
                file_count="multiple",
            )
            upload_btn = gr.Button("Process & Add", variant="primary")
            status_box = gr.Textbox(
                label="Status",
                lines=5,
                interactive=False,
                placeholder="Upload files and click Process & Add...",
            )

            gr.Markdown("---")

            gr.Markdown("### Loaded documents")
            doc_list_box = gr.Textbox(
                label="Documents in memory",
                lines=5,
                interactive=False,
                value="No documents loaded.",
            )
            delete_dropdown = gr.Dropdown(
                label="Select document to delete",
                choices=[],
                value=None,
                interactive=True,
            )
            with gr.Row():
                delete_btn     = gr.Button("Delete selected", variant="secondary")
                delete_all_btn = gr.Button("Delete all", variant="stop")

            hard_reset_btn = gr.Button("Hard Reset (clear all)", variant="stop")

        # ── Right panel ─────────────────────────────────────────────────────
        with gr.Column(scale=2):

            gr.Markdown("### Chat")
            chatbot = gr.Chatbot(
                label="Conversation",
                height=480,
            )
            with gr.Row():
                msg_input = gr.Textbox(
                    label="Your question",
                    placeholder="Ask something about your documents...",
                    scale=4,
                    lines=1,
                )
                send_btn = gr.Button("Send", variant="primary", scale=1)
            clear_btn = gr.Button("Clear chat")

    # ── Wire up events ────────────────────────────────────────────────────────

    upload_btn.click(
        fn=upload_files,
        inputs=[file_input],
        outputs=[status_box, doc_list_box, delete_dropdown],
    )

    delete_btn.click(
        fn=delete_document,
        inputs=[delete_dropdown],
        outputs=[status_box, doc_list_box, delete_dropdown],
    )

    delete_all_btn.click(
        fn=delete_all_documents,
        outputs=[status_box, doc_list_box, delete_dropdown, chatbot, msg_input],
    )

    hard_reset_btn.click(
        fn=hard_reset,
        outputs=[status_box, doc_list_box, delete_dropdown, chatbot, msg_input],
    )

    send_btn.click(
        fn=chat,
        inputs=[msg_input, chatbot],
        outputs=[msg_input, chatbot],
    )

    msg_input.submit(
        fn=chat,
        inputs=[msg_input, chatbot],
        outputs=[msg_input, chatbot],
    )

    clear_btn.click(
        fn=clear_chat,
        outputs=[chatbot, msg_input],
    )


if __name__ == "__main__":
    try:
        demo.launch(
            server_name="0.0.0.0",
            server_port=7860,
            share=True,
            theme=gr.themes.Soft(),
        )
    except KeyboardInterrupt:
        print("\n👋 App stopped.")
