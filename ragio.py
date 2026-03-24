"""
RAG Gradio App — LangChain + Ollama + PDF/DOCX
Full fixed version — clear chat + error handling
-------------------------------------------------
Requirements:
    pip install gradio langchain langchain-community langchain-ollama
                langchain-text-splitters langchain-chroma langchain-core
                chromadb pypdf python-docx docx2txt

Run:
    python rag_gradio.py
Then open: http://localhost:7860
"""
import gc
import os
import time
import shutil
import gradio as gr
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_ollama import OllamaEmbeddings, ChatOllama
from langchain_core.prompts import PromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser


# ── Config ────────────────────────────────────────────────────────────────────

OLLAMA_MODEL  = "llama3"
EMBED_MODEL   = "nomic-embed-text"
CHROMA_DIR    = "./chroma_db"
CHUNK_SIZE    = 500
CHUNK_OVERLAP = 100
RETRIEVER_K   = 6
ALLOWED_EXTS  = {".pdf", ".docx"}

PROMPT_TEMPLATE = """### SYSTEM:
You are a strict document assistant. Your task is to answer the question using ONLY the provided context. 

### CONSTRAINTS:
1. Use ONLY the information in the 'Context' section.
2. If the answer is not explicitly stated in the context, you MUST say: "I am sorry, but the uploaded documents do not contain information regarding this request."
3. Do NOT use any outside knowledge or provide explanations not found in the text.
4. Do NOT mention the context or documents in your final answer (e.g., don't say "According to the context...").

### Context:
{context}

### Question: 
{question}

### Answer:"""

# ── Global state ──────────────────────────────────────────────────────────────

vectorstore = None
chain       = None
doc_store   = {}


# ── Helpers ───────────────────────────────────────────────────────────────────

def validate_files(files):
    if not files:
        return [], [("none", "no file")]
    invalid = []
    valid   = []
    for file in files:
        fname = os.path.basename(file.name)
        ext   = os.path.splitext(fname)[-1].lower()
        if ext not in ALLOWED_EXTS:
            invalid.append((fname, ext if ext else "no extension"))
        else:
            valid.append(file)
    return valid, invalid


def load_document(path: str):
    ext = os.path.splitext(path)[-1].lower()
    if ext == ".pdf":
        loader = PyPDFLoader(path)
        pages  = loader.load()
    elif ext == ".docx":
        docx    = DocxDocument(path)
        content = "\n".join([p.text for p in docx.paragraphs if p.text.strip()])
        pages   = [Document(
            page_content=content,
            metadata={"filename": os.path.basename(path), "page": 0},
        )]
    else:
        raise ValueError(f"Unsupported file type: {ext}")
    for i, page in enumerate(pages):
        page.metadata.setdefault("page", i)
        page.metadata["filename"] = os.path.basename(path)
        page.metadata["filetype"] = ext
    return pages


def rebuild_vectorstore():
	global vectorstore, chain

	# Step 1: Force release of all LangChain and Chroma objects
	chain = None
	if vectorstore is not None:
		try:
			# Closes the underlying PersistentClient
			vectorstore._client.close()
		except Exception:
			pass
		vectorstore = None

	# Step 2: Trigger Garbage Collection
	# This is critical on Windows to release file locks immediately
	gc.collect()
	time.sleep(1)  # Brief pause for the OS to finalize the release

	# Step 3: Attempt folder deletion with retries
	if os.path.exists(CHROMA_DIR):
		for attempt in range(5):
			try:
				shutil.rmtree(CHROMA_DIR)
				print(f"Successfully cleared old vector store on attempt {attempt + 1}")
				break
			except PermissionError:
				time.sleep(1)
		else:
			# If all 5 attempts fail
			raise Exception(
				"Windows is still locking the database files. "
				"Please close any other programs using this folder and try again."
			)

	# Step 4: Rebuild only if there are documents in the store
	all_pages = []
	for pages in doc_store.values():
		all_pages.extend(pages)

	if not all_pages:
		return

	splitter = RecursiveCharacterTextSplitter(
		chunk_size=CHUNK_SIZE,
		chunk_overlap=CHUNK_OVERLAP,
	)
	chunks = splitter.split_documents(all_pages)

	# Re-initialize with a fresh client
	embeddings = OllamaEmbeddings(model=EMBED_MODEL)
	vectorstore = Chroma.from_documents(
		documents=chunks,
		embedding=embeddings,
		persist_directory=CHROMA_DIR,
	)

	# Rebuild the chain with the new vectorstore
	chain = build_chain(vectorstore)
	print(f"✅ Vector store rebuilt with {len(chunks)} chunks.")
def hard_reset():
    global vectorstore, chain, doc_store

    # 1. Break references
    chain = None
    if vectorstore is not None:
        try:
            vectorstore._client.close()
        except Exception:
            pass
    vectorstore = None
    doc_store = {}

    # 2. Flush memory
    gc.collect()
    time.sleep(1)

    # 3. Wipe physical directory
    if os.path.exists(CHROMA_DIR):
        for _ in range(5):
            try:
                shutil.rmtree(CHROMA_DIR)
                break
            except PermissionError:
                time.sleep(1)

    return (
        "🔄 Hard reset complete — all documents and vectors cleared.",
        "No documents loaded.",
        gr.update(choices=[], value=None),
        [], # Clears chatbot history
        "", # Clears message input
    )
def build_chain(vs):
	prompt = PromptTemplate.from_template(PROMPT_TEMPLATE)
	llm = ChatOllama(model=OLLAMA_MODEL, temperature=0)  # Temperature 0 is correct
	retriever = vs.as_retriever(
		search_type="similarity",
		search_kwargs={"k": 6},  # Increased k
	)

	def format_docs(docs):
		# Adding metadata (filename) helps the model distinguish between sources
		formatted = []
		for d in docs:
			fname = d.metadata.get("filename", "Unknown")
			formatted.append(f"[Source: {fname}]\n{d.page_content}")
		return "\n\n---\n\n".join(formatted)

	return (
			{"context": retriever | format_docs, "question": RunnablePassthrough()}
			| prompt
			| llm
			| StrOutputParser()
	)

def format_doc_list():
    if not doc_store:
        return "No documents loaded."
    lines = []
    for fname, pages in doc_store.items():
        ext = os.path.splitext(fname)[-1].upper().replace(".", "")
        lines.append(f"• {fname}  [{ext} — {len(pages)} section(s)]")
    return "\n".join(lines)


def get_doc_choices():
    return list(doc_store.keys())


# ── Gradio handlers ───────────────────────────────────────────────────────────

def upload_files(files):
    global doc_store
    # Safety check — warn if old chroma still exists
    if os.path.exists(CHROMA_DIR) and not doc_store:
        shutil.rmtree(CHROMA_DIR, ignore_errors=True)
		
    if not files:
        raise gr.Error("No files selected. Please upload at least one .pdf or .docx file.")

    valid_files, invalid = validate_files(files)

    if invalid:
        rejected = ", ".join(f"{fname} ('{ext}')" for fname, ext in invalid)
        raise gr.Error(
            f"Unsupported file type(s): {rejected}. "
            f"Only .pdf and .docx files are accepted. "
            f"Please remove the invalid file(s) and try again."
        )

    newly_added = []
    load_errors = []

    for file in valid_files:
        fname = os.path.basename(file.name)
        if fname in doc_store:
            load_errors.append(f"⚠️ '{fname}' already loaded — skipped.")
            continue
        try:
            pages            = load_document(file.name)
            doc_store[fname] = pages
            newly_added.append(fname)
        except Exception as e:
            load_errors.append(f"❌ Failed to load '{fname}': {e}")

    if newly_added:
        try:
            rebuild_vectorstore()
        except Exception as e:
            raise gr.Error(f"Error building vector store: {e}")

    parts = []
    if newly_added:
        parts.append("✅ Successfully added:\n" + "\n".join(f"  • {f}" for f in newly_added))
    if load_errors:
        parts.append("\n".join(load_errors))
    parts.append(f"\n📦 Total loaded: {len(doc_store)} file(s)")

    return (
        "\n".join(parts),
        format_doc_list(),
        gr.update(choices=get_doc_choices(), value=None),
    )


def delete_document(filename):
    global doc_store

    if not filename:
        raise gr.Error("Please select a document from the dropdown first.")

    if filename not in doc_store:
        raise gr.Error(f"'{filename}' was not found in loaded documents.")

    del doc_store[filename]
    rebuild_vectorstore()

    status = f"🗑️ Deleted: {filename}\n"
    status += (
        f"📦 Remaining: {len(doc_store)} file(s)"
        if doc_store
        else "📭 No documents loaded. Please upload new files."
    )

    return (
        status,
        format_doc_list(),
        gr.update(choices=get_doc_choices(), value=None),
    )


def delete_all_documents():
    global doc_store, vectorstore, chain

    if vectorstore is not None:
        try:
            vectorstore._client.close()
        except Exception:
            pass

    doc_store   = {}
    vectorstore = None
    chain       = None

    if os.path.exists(CHROMA_DIR):
        for attempt in range(3):
            try:
                shutil.rmtree(CHROMA_DIR)
                break
            except PermissionError:
                time.sleep(1)

    return (
        "🗑️ All documents deleted. Please upload new files.",
        format_doc_list(),
        gr.update(choices=[], value=None),
        [],
        "",
    )


def chat(message, history):
    global chain

    history = history or []

    if not message.strip():
        return "", history

    if chain is None:
        history.append({"role": "user",      "content": message})
        history.append({"role": "assistant", "content": "⚠️ No documents loaded. Please upload a .pdf or .docx file first."})
        return "", history

    try:
        answer = chain.invoke(message)
        history.append({"role": "user",      "content": message})
        history.append({"role": "assistant", "content": answer})
    except Exception as e:
        history.append({"role": "user",      "content": message})
        history.append({"role": "assistant", "content": f"⚠️ Error: {str(e)}. Please try again."})

    return "", history


def clear_chat():
    # Returns empty history and clears input box
    return [], ""


# ── Gradio UI ─────────────────────────────────────────────────────────────────

with gr.Blocks(title="RAG Chat — Ollama") as demo:

    gr.Markdown("# Documents Summarizer")
    gr.Markdown("Upload PDF or DOCX files, manage your document library, then chat.")

    hard_reset_btn = gr.Button("Hard Reset (clear all)", variant="stop")

	# # wire it up:
    # hard_reset_btn.click(
	# 	fn=hard_reset,
	# 	outputs=[status_box, doc_list_box, delete_dropdown, chatbot, msg_input],
	# )
    with gr.Row():

        # ── Left panel ──────────────────────────────────────────────────────
        with gr.Column(scale=1):

            gr.Markdown("### Upload documents")
            gr.Markdown(
                "<small style='color:gray'>Accepted: <b>.pdf</b> and <b>.docx</b> only. "
                "Other file types will show an error.</small>"
            )
            file_input = gr.File(
                label="PDF / DOCX files",
                file_types=[".pdf", ".docx"],
                file_count="multiple",
            )
            upload_btn = gr.Button("Process & Add", variant="primary")
            status_box = gr.Textbox(
                label="Status",
                lines=5,
                interactive=False,
                placeholder="Upload files and click Process & Add...",
            )

            gr.Markdown("---")

            gr.Markdown("### Loaded documents")
            doc_list_box = gr.Textbox(
                label="Documents in memory",
                lines=5,
                interactive=False,
                value="No documents loaded.",
            )
            delete_dropdown = gr.Dropdown(
                label="Select document to delete",
                choices=[],
                value=None,
                interactive=True,
            )
            with gr.Row():
                delete_btn     = gr.Button("Delete selected", variant="secondary")
                delete_all_btn = gr.Button("Delete all", variant="stop")

        # ── Right panel ─────────────────────────────────────────────────────
        with gr.Column(scale=2):

            gr.Markdown("### Chat")
            chatbot = gr.Chatbot(
                label="Conversation",
                height=480,
            )
            with gr.Row():
                msg_input = gr.Textbox(
                    label="Your question",
                    placeholder="Ask something about your documents...",
                    scale=4,
                    lines=1,
                )
                send_btn = gr.Button("Send", variant="primary", scale=1)
				# ✅ PLACE IT HERE (After all UI elements like status_box are defined)
                hard_reset_btn.click(
					fn=hard_reset,
					outputs=[status_box, doc_list_box, delete_dropdown, chatbot, msg_input],
				)
            # ── Clear chat button — inside gr.Blocks so variables are in scope
            clear_btn = gr.Button("Clear chat")

    # ── Wire up events (all inside with gr.Blocks) ────────────────────────────

    upload_btn.click(
        fn=upload_files,
        inputs=[file_input],
        outputs=[status_box, doc_list_box, delete_dropdown],
    )

    delete_btn.click(
        fn=delete_document,
        inputs=[delete_dropdown],
        outputs=[status_box, doc_list_box, delete_dropdown],
    )

    delete_all_btn.click(
        fn=delete_all_documents,
        outputs=[status_box, doc_list_box, delete_dropdown, chatbot, msg_input],
    )

    send_btn.click(
        fn=chat,
        inputs=[msg_input, chatbot],
        outputs=[msg_input, chatbot],
    )

    msg_input.submit(
        fn=chat,
        inputs=[msg_input, chatbot],
        outputs=[msg_input, chatbot],
    )

    # ── Clear chat: wipes both chatbot history and input box ──────────────────
    clear_btn.click(
        fn=clear_chat,
        outputs=[chatbot, msg_input],   # ← both cleared at once
    )


if __name__ == "__main__":
    try:
        demo.launch(
            server_name="0.0.0.0",
            server_port=7860,
            share=True,
            theme=gr.themes.Soft(),
        )
    except KeyboardInterrupt:
        print("\n👋 App stopped.")
